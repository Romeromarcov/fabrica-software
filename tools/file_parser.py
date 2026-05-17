"""
Parsea archivos subidos por el Founder para dar contexto al Agente 0.

Formatos soportados: .txt .md .json .yaml .yml .csv .pdf .docx
"""
import json
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".txt", ".md", ".json", ".yaml", ".yml", ".csv", ".pdf", ".docx"}

# Límite por archivo — suficiente para un PRD completo sin explotar el contexto
MAX_CHARS_PER_FILE = 10_000


def parse_file(path: Path) -> str:
    """Lee un archivo y devuelve su contenido como texto plano truncado."""
    ext = path.suffix.lower()

    try:
        if ext in (".txt", ".md", ".yaml", ".yml"):
            return path.read_text(encoding="utf-8", errors="replace")[:MAX_CHARS_PER_FILE]

        if ext == ".json":
            raw = path.read_text(encoding="utf-8", errors="replace")
            try:
                data = json.loads(raw)
                return json.dumps(data, indent=2, ensure_ascii=False)[:MAX_CHARS_PER_FILE]
            except json.JSONDecodeError:
                return raw[:MAX_CHARS_PER_FILE]

        if ext == ".csv":
            lines = path.read_text(encoding="utf-8", errors="replace").splitlines()
            preview = "\n".join(lines[:60])
            tail = f"\n\n[... {len(lines)-60} filas adicionales omitidas]" if len(lines) > 60 else ""
            return (preview + tail)[:MAX_CHARS_PER_FILE]

        if ext == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(path))
                pages = []
                for i, page in enumerate(reader.pages[:25]):
                    text = page.extract_text() or ""
                    if text.strip():
                        pages.append(f"[Página {i+1}]\n{text}")
                return "\n\n".join(pages)[:MAX_CHARS_PER_FILE]
            except ImportError:
                return "⚠️  PDF no procesado — ejecuta: pip install pypdf"
            except Exception as e:
                logger.warning("Error leyendo PDF %s: %s", path.name, e)
                return f"⚠️  Error leyendo PDF: {e}"

        if ext == ".docx":
            try:
                from docx import Document
                doc = Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                # Incluir también texto de tablas
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text:
                            paragraphs.append(row_text)
                return "\n\n".join(paragraphs)[:MAX_CHARS_PER_FILE]
            except ImportError:
                return "⚠️  DOCX no procesado — ejecuta: pip install python-docx"
            except Exception as e:
                logger.warning("Error leyendo DOCX %s: %s", path.name, e)
                return f"⚠️  Error leyendo DOCX: {e}"

    except Exception as e:
        logger.error("Error inesperado parseando %s: %s", path.name, e)
        return f"⚠️  Error leyendo archivo: {e}"

    return f"⚠️  Formato no soportado: {ext}"


def read_uploads(project_id: str) -> dict[str, str]:
    """
    Lee todos los archivos del directorio uploads/ de un proyecto.
    Retorna {nombre_archivo: contenido_texto}.
    """
    from config import RUNS_DIR
    uploads_dir = RUNS_DIR / project_id / "uploads"
    if not uploads_dir.exists():
        return {}

    result: dict[str, str] = {}
    for f in sorted(uploads_dir.iterdir()):
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS:
            content = parse_file(f)
            if content.strip():
                result[f.name] = content
                logger.info("Archivo parseado: %s (%d chars)", f.name, len(content))

    return result


def uploads_as_context(project_id: str) -> str:
    """
    Convierte los archivos subidos en un bloque de contexto listo para
    inyectar en el prompt del Agente 0.
    Devuelve cadena vacía si no hay archivos.
    """
    uploads = read_uploads(project_id)
    if not uploads:
        return ""

    sections = []
    for filename, content in uploads.items():
        ext  = Path(filename).suffix.lower()
        icon = {"pdf": "📄", "docx": "📝", "md": "📋", "json": "🔧",
                "csv": "📊", "yaml": "⚙️", "yml": "⚙️"}.get(ext.lstrip("."), "📎")
        sections.append(f"### {icon} {filename}\n\n{content}")

    header = (
        f"## DOCUMENTOS APORTADOS POR EL FOUNDER ({len(uploads)} archivo{'s' if len(uploads)>1 else ''})\n\n"
        "Estos documentos contienen la visión, requisitos o contexto que el Founder "
        "proporcionó para el proyecto. Tienen prioridad sobre cualquier suposición.\n\n"
    )
    return header + "\n\n---\n\n".join(sections)
